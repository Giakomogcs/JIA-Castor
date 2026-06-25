#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Conversor Markdown -> DOCX com identidade visual SENAI para a documentacao Castor.
Renderiza: capa, header/footer, TOC, titulos coloridos, tabelas zebradas,
callouts, passos numerados, blocos de codigo, diagramas Mermaid (via mmdc) e
markdown inline (**negrito**, *italico*, `codigo`).
"""
import os
import re
import shutil
import subprocess
import tempfile

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = os.path.dirname(os.path.abspath(__file__))
MD_PATH = os.path.join(BASE, "DOCUMENTACAO.md")
DOCX_PATH = os.path.join(BASE, "Castor-Documentacao.docx")

# ---- Paleta SENAI ----
ACCENT = "E30613"
ACCENT2 = "8B0410"
INK = "0B0B0B"
INK2 = "1F2937"
GRAY = "6B7280"
RULE = "D1D5DB"
ZEBRA = "F9F7F7"
CELLHDR = "FBE9EB"
CODE_BG = "0B0B0B"
CODE_FG = "E5E7EB"
CALLOUTS = {
    "NOTE": ("1D4ED8", "EFF4FF", "i", "Nota"),
    "INFO": ("1D4ED8", "EFF4FF", "i", "Informacao"),
    "TIP": ("047857", "ECFDF5", "+", "Dica"),
    "WARNING": ("B45309", "FFFBEB", "!", "Atencao"),
    "DANGER": ("B91C1C", "FEF2F2", "x", "Critico"),
}
FONT = "Calibri"
MONO = "Consolas"


def rgb(hexstr):
    return RGBColor.from_string(hexstr)


def set_cell_bg(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hexcolor)
    tcPr.append(shd)


def set_cell_borders(cell, color=RULE, sz=4, sides=("top", "left", "bottom", "right")):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for side in sides:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(sz))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        borders.append(el)
    tcPr.append(borders)


def set_cell_margins(cell, top=60, bottom=60, left=110, right=110):
    tcPr = cell._tc.get_or_add_tcPr()
    m = OxmlElement("w:tcMar")
    for side, val in (("top", top), ("bottom", bottom), ("start", left), ("end", right)):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"), str(val))
        el.set(qn("w:type"), "dxa")
        m.append(el)
    tcPr.append(m)


def shade_paragraph(p, hexcolor):
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hexcolor)
    pPr.append(shd)


def left_bar(p, hexcolor, sz=24):
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), str(sz))
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), hexcolor)
    pbdr.append(left)
    pPr.append(pbdr)


def bottom_rule(p, hexcolor, sz=6):
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(sz))
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), hexcolor)
    pbdr.append(bottom)
    pPr.append(pbdr)


INLINE_RE = re.compile(r"(\*\*.+?\*\*|\*.+?\*|`.+?`)")


def add_inline(paragraph, text, base_color=INK2, base_size=10.5, bold=False):
    text = text.replace("&amp;", "&").replace("&lt;", "<").replace("&gt;", ">")
    parts = INLINE_RE.split(text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            r = paragraph.add_run(part[2:-2])
            r.bold = True
        elif part.startswith("`") and part.endswith("`"):
            r = paragraph.add_run(part[1:-1])
            r.font.name = MONO
            r.font.color.rgb = rgb(ACCENT2)
            r.font.size = Pt(base_size - 0.5)
        elif part.startswith("*") and part.endswith("*"):
            r = paragraph.add_run(part[1:-1])
            r.italic = True
        else:
            r = paragraph.add_run(part)
        r.font.name = FONT
        r.font.color.rgb = rgb(base_color)
        r.font.size = Pt(base_size)
        if bold:
            r.bold = True


# ---------- Field helpers (TOC + page number) ----------
def add_field(paragraph, instr):
    run = paragraph.add_run()
    fld = OxmlElement("w:fldChar")
    fld.set(qn("w:fldCharType"), "begin")
    run._r.append(fld)
    run2 = paragraph.add_run()
    it = OxmlElement("w:instrText")
    it.set(qn("xml:space"), "preserve")
    it.text = instr
    run2._r.append(it)
    run3 = paragraph.add_run()
    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    run3._r.append(sep)
    run4 = paragraph.add_run()
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run4._r.append(end)
    return run4


def add_toc(document):
    p = document.add_paragraph()
    add_field(p, 'TOC \\o "1-3" \\h \\z \\u')


# ---------- Mermaid ----------
def render_mermaid(code, idx):
    mmdc = shutil.which("mmdc")
    if not mmdc:
        return None
    tmpdir = tempfile.mkdtemp()
    mmd = os.path.join(tmpdir, f"d{idx}.mmd")
    png = os.path.join(BASE, f"_diagram_{idx}.png")
    with open(mmd, "w", encoding="utf-8") as f:
        f.write(code)
    cfg = os.path.join(tmpdir, "cfg.json")
    with open(cfg, "w", encoding="utf-8") as f:
        f.write('{"theme":"base","themeVariables":{"primaryColor":"#FBE9EB",'
                '"primaryBorderColor":"#E30613","primaryTextColor":"#0B0B0B",'
                '"lineColor":"#8B0410","fontFamily":"Calibri"}}')
    try:
        subprocess.run(
            [mmdc, "-i", mmd, "-o", png, "-b", "white", "-c", cfg, "-s", "2"],
            check=True, capture_output=True, timeout=120, shell=True,
        )
        if os.path.exists(png):
            return png
    except Exception as e:
        print(f"  [mermaid] falha no diagrama {idx}: {e}")
    return None


# ---------- Table parsing ----------
def is_table_sep(line):
    return bool(re.match(r"^\s*\|?[\s:|-]+\|?\s*$", line)) and "-" in line


def split_row(line):
    line = line.strip()
    if line.startswith("|"):
        line = line[1:]
    if line.endswith("|"):
        line = line[:-1]
    return [c.strip() for c in line.split("|")]


def add_table(document, header, rows):
    ncols = len(header)
    table = document.add_table(rows=1, cols=ncols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    hdr = table.rows[0].cells
    for i, htext in enumerate(header):
        set_cell_bg(hdr[i], ACCENT)
        set_cell_borders(hdr[i], RULE)
        set_cell_margins(hdr[i])
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(htext.replace("**", ""))
        r.bold = True
        r.font.name = FONT
        r.font.size = Pt(9.5)
        r.font.color.rgb = rgb("FFFFFF")
    for ridx, row in enumerate(rows):
        cells = table.add_row().cells
        for i in range(ncols):
            val = row[i] if i < len(row) else ""
            set_cell_borders(cells[i], RULE)
            set_cell_margins(cells[i])
            if ridx % 2 == 1:
                set_cell_bg(cells[i], ZEBRA)
            p = cells[i].paragraphs[0]
            add_inline(p, val, base_color=INK2, base_size=9.5)
    document.add_paragraph()


# ---------- Callout ----------
def add_callout(document, kind, lines):
    color, bg, icon, label = CALLOUTS.get(kind, CALLOUTS["NOTE"])
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    set_cell_bg(cell, bg)
    set_cell_borders(cell, color, sz=4, sides=("top", "bottom", "right"))
    # thick left bar
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    el = OxmlElement("w:left")
    el.set(qn("w:val"), "single")
    el.set(qn("w:sz"), "24")
    el.set(qn("w:space"), "0")
    el.set(qn("w:color"), color)
    borders.append(el)
    tcPr.append(borders)
    set_cell_margins(cell, top=80, bottom=80, left=160, right=140)
    p = cell.paragraphs[0]
    r = p.add_run(f"[{label.upper()}]  ")
    r.bold = True
    r.font.name = FONT
    r.font.size = Pt(10)
    r.font.color.rgb = rgb(color)
    add_inline(p, " ".join(lines), base_color=INK2, base_size=10)
    document.add_paragraph()


# ---------- Code block ----------
def add_code_block(document, lines):
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    set_cell_bg(cell, CODE_BG)
    set_cell_borders(cell, CODE_BG)
    set_cell_margins(cell, top=120, bottom=120, left=160, right=160)
    first = True
    for ln in lines:
        p = cell.paragraphs[0] if first else cell.add_paragraph()
        first = False
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        r = p.add_run(ln if ln else " ")
        r.font.name = MONO
        r.font.size = Pt(9)
        r.font.color.rgb = rgb(CODE_FG)
    document.add_paragraph()


# ---------- Headings ----------
def add_heading(document, level, text):
    text = text.replace("**", "")
    if level == 1:
        p = document.add_paragraph(style="Heading 1")
        r = p.add_run(text.upper())
        r.font.name = FONT
        r.font.size = Pt(18)
        r.bold = True
        r.font.color.rgb = rgb(ACCENT)
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(6)
        bottom_rule(p, ACCENT, sz=12)
    elif level == 2:
        p = document.add_paragraph(style="Heading 2")
        r = p.add_run(text)
        r.font.name = FONT
        r.font.size = Pt(14)
        r.bold = True
        r.font.color.rgb = rgb(INK)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        bottom_rule(p, RULE, sz=6)
    else:
        p = document.add_paragraph(style="Heading 3")
        r = p.add_run(text)
        r.font.name = FONT
        r.font.size = Pt(11.5)
        r.bold = True
        r.font.color.rgb = rgb(INK2)
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(2)


# ---------- Numbered step ----------
def add_step(document, num, text):
    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(f" {num} ")
    r.bold = True
    r.font.name = FONT
    r.font.size = Pt(10)
    r.font.color.rgb = rgb("FFFFFF")
    # red badge background on the run
    rpr = r._r.get_or_add_rPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), ACCENT)
    rpr.append(shd)
    p.add_run("  ")
    add_inline(p, text, base_color=INK2, base_size=10.5)


def add_bullet(document, text):
    p = document.add_paragraph(style="List Bullet")
    add_inline(p, text, base_color=INK2, base_size=10.5)


def add_para(document, text):
    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    add_inline(p, text, base_color=INK2, base_size=10.5)


# ---------- Header / Footer ----------
def setup_header_footer(section, doc_title, project):
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.text = ""
    # tab stops: left title, right project
    pPr = hp._p.get_or_add_pPr()
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "right")
    tab.set(qn("w:pos"), "9360")
    tabs.append(tab)
    pPr.append(tabs)
    r = hp.add_run(doc_title)
    r.font.name = FONT
    r.font.size = Pt(8)
    r.font.color.rgb = rgb(GRAY)
    r2 = hp.add_run("\t" + project)
    r2.font.name = FONT
    r2.font.size = Pt(8)
    r2.font.color.rgb = rgb(ACCENT)
    r2.bold = True
    bottom_rule(hp, RULE, sz=4)

    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.text = ""
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = fp.add_run("Pagina ")
    r.font.name = FONT
    r.font.size = Pt(8)
    r.font.color.rgb = rgb(GRAY)
    add_field(fp, "PAGE")
    r2 = fp.add_run(" de ")
    r2.font.name = FONT
    r2.font.size = Pt(8)
    r2.font.color.rgb = rgb(GRAY)
    add_field(fp, "NUMPAGES")
    for run in fp.runs:
        run.font.name = FONT
        run.font.size = Pt(8)
        run.font.color.rgb = rgb(GRAY)


# ---------- Cover ----------
def build_cover(document, title, subtitle, project, date_str, author):
    # red band
    t = document.add_table(rows=1, cols=1)
    cell = t.rows[0].cells[0]
    set_cell_bg(cell, ACCENT)
    set_cell_borders(cell, ACCENT)
    set_cell_margins(cell, top=400, bottom=400, left=300, right=300)
    p = cell.paragraphs[0]
    r = p.add_run("SENAI  ·  PROJETO DE IA NA INDUSTRIA")
    r.font.name = FONT
    r.font.size = Pt(11)
    r.bold = True
    r.font.color.rgb = rgb("FFFFFF")

    document.add_paragraph()
    document.add_paragraph()

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(title)
    r.font.name = FONT
    r.font.size = Pt(40)
    r.bold = True
    r.font.color.rgb = rgb(ACCENT)

    p = document.add_paragraph()
    r = p.add_run(subtitle)
    r.font.name = FONT
    r.font.size = Pt(16)
    r.font.color.rgb = rgb(INK2)

    document.add_paragraph()
    p = document.add_paragraph()
    bottom_rule(p, ACCENT2, sz=12)

    document.add_paragraph()
    for label, value in [("Projeto", project), ("Data", date_str), ("Autor", author)]:
        p = document.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(f"{label}:  ")
        r.font.name = FONT
        r.font.size = Pt(11)
        r.bold = True
        r.font.color.rgb = rgb(ACCENT2)
        r2 = p.add_run(value)
        r2.font.name = FONT
        r2.font.size = Pt(11)
        r2.font.color.rgb = rgb(INK2)

    document.add_page_break()


def main():
    with open(MD_PATH, "r", encoding="utf-8") as f:
        raw = f.read()
    lines = raw.split("\n")

    document = Document()
    # base style
    normal = document.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = rgb(INK2)

    section = document.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    DOC_TITLE = "Castor — Copiloto Comercial B2B"
    PROJECT = "Castor · SENAI"

    build_cover(
        document,
        "Castor",
        "Copiloto Comercial B2B — Documentacao de Negocio e Tecnica",
        "Castor (distribuidora) · IA na Industria / SENAI",
        "24/06/2026",
        "Doc Master",
    )

    setup_header_footer(section, DOC_TITLE, PROJECT)

    # TOC
    add_heading(document, 1, "Sumario")
    add_toc(document)
    document.add_page_break()

    diagram_count = 0
    i = 0
    n = len(lines)
    while i < n:
        line = lines[i]
        stripped = line.strip()

        # skip the top H1 title + first lines already on cover
        if i < 6 and (stripped.startswith("# Castor") or stripped == "Documentação Completa de Negócio e Técnica"
                      or stripped == "Projeto de IA na Indústria — Parceria SENAI" or stripped == "---" or stripped == ""):
            i += 1
            continue

        if not stripped:
            i += 1
            continue

        if stripped == "---":
            i += 1
            continue

        # code / mermaid fence
        if stripped.startswith("```"):
            lang = stripped[3:].strip()
            block = []
            i += 1
            while i < n and not lines[i].strip().startswith("```"):
                block.append(lines[i])
                i += 1
            i += 1  # skip closing fence
            if lang == "mermaid":
                diagram_count += 1
                png = render_mermaid("\n".join(block), diagram_count)
                if png:
                    p = document.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run()
                    # scale to max ~6.2 inches wide
                    run.add_picture(png, width=Inches(6.0))
                    cap = document.add_paragraph()
                    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    cr = cap.add_run(f"Figura {diagram_count} — diagrama")
                    cr.italic = True
                    cr.font.name = FONT
                    cr.font.size = Pt(8.5)
                    cr.font.color.rgb = rgb(GRAY)
                    try:
                        os.remove(png)
                    except OSError:
                        pass
                else:
                    add_code_block(document, block + ["", "(diagrama Mermaid)"])
            else:
                add_code_block(document, block)
            continue

        # callout
        m = re.match(r"^>\s*\[!(\w+)\]", stripped)
        if m:
            kind = m.group(1).upper()
            clines = []
            i += 1
            while i < n and lines[i].strip().startswith(">"):
                clines.append(lines[i].strip().lstrip(">").strip())
                i += 1
            add_callout(document, kind, [c for c in clines if c])
            continue

        # blockquote (plain)
        if stripped.startswith(">"):
            qlines = []
            while i < n and lines[i].strip().startswith(">"):
                qlines.append(lines[i].strip().lstrip(">").strip())
                i += 1
            p = document.add_paragraph()
            shade_paragraph(p, ZEBRA)
            left_bar(p, GRAY, sz=18)
            p.paragraph_format.space_after = Pt(6)
            add_inline(p, " ".join([q for q in qlines if q]), base_color=GRAY, base_size=10)
            continue

        # headings
        if stripped.startswith("#"):
            hm = re.match(r"^(#+)\s+(.*)$", stripped)
            level = len(hm.group(1))
            text = hm.group(2)
            if text.strip().lower() == "sumário executivo":
                add_heading(document, 1, text)
            else:
                add_heading(document, min(level, 3) if level > 1 else 1, text)
            i += 1
            continue

        # table
        if stripped.startswith("|") and i + 1 < n and is_table_sep(lines[i + 1]):
            header = split_row(lines[i])
            i += 2
            rows = []
            while i < n and lines[i].strip().startswith("|"):
                rows.append(split_row(lines[i]))
                i += 1
            add_table(document, header, rows)
            continue

        # numbered step
        nm = re.match(r"^(\d+)\.\s+(.*)$", stripped)
        if nm:
            add_step(document, nm.group(1), nm.group(2))
            i += 1
            continue

        # bullet
        if stripped.startswith("- ") or stripped.startswith("* "):
            add_bullet(document, stripped[2:])
            i += 1
            continue

        # italic-only footer line
        if stripped.startswith("*") and stripped.endswith("*") and not stripped.startswith("**"):
            p = document.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(stripped.strip("*"))
            r.italic = True
            r.font.name = FONT
            r.font.size = Pt(9)
            r.font.color.rgb = rgb(GRAY)
            i += 1
            continue

        # paragraph
        add_para(document, stripped)
        i += 1

    document.save(DOCX_PATH)
    print(f"OK: {DOCX_PATH}")


if __name__ == "__main__":
    main()
